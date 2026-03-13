#!/usr/bin/env python3
"""
加密货币新闻文章收集器
功能：从多个加密货币新闻源收集最新热门文章，翻译成中文并保存为Word文档
"""

import requests
from bs4 import BeautifulSoup
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time
import json
import re
import os
import sys
import shutil
import random
import hashlib
import feedparser

# ========================= 配置 =========================
TOP_N = 3               # 最终选取的文章数量
CANDIDATE_N = 15         # 候选文章数量（从中筛选 TOP_N 篇）
REQUEST_TIMEOUT = 20     # 网络请求超时（秒）
TRANSLATE_CHUNK = 4500   # 翻译分块大小（字符数）
RETRY_TIMES = 3          # 网络请求重试次数
REQUEST_DELAY = 2        # 请求间隔（秒）
# ========================================================


class NetworkHelper:
    """网络请求辅助类，提供重试和随机 UA"""

    USER_AGENTS = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.3 Safari/605.1.15',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:123.0) Gecko/20100101 Firefox/123.0',
        'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36',
    ]

    @classmethod
    def get_headers(cls, accept='text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'):
        return {
            'User-Agent': random.choice(cls.USER_AGENTS),
            'Accept': accept,
            'Accept-Language': 'en-US,en;q=0.9',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
        }

    @classmethod
    def get(cls, url, timeout=REQUEST_TIMEOUT, retries=RETRY_TIMES, **kwargs):
        """带重试的 GET 请求"""
        # 从 kwargs 提取 headers（只提取一次，避免重试时丢失）
        custom_headers = kwargs.pop('headers', None)
        last_err = None
        for attempt in range(retries):
            try:
                headers = custom_headers or cls.get_headers()
                # 使用 (connect_timeout, read_timeout) 元组，连接超时更短
                t = (min(timeout, 10), timeout)
                resp = requests.get(url, headers=headers, timeout=t, **kwargs)
                resp.raise_for_status()
                return resp
            except KeyboardInterrupt:
                raise
            except Exception as e:
                last_err = e
                if attempt < retries - 1:
                    wait = (attempt + 1) * 2
                    print(f"  请求失败 (重试 {attempt + 1}/{retries}): {e}")
                    time.sleep(wait)
        raise last_err


class CoinDeskCollector:
    """从多个加密货币新闻 RSS 源收集文章"""

    def __init__(self):
        self.base_url = "https://www.coindesk.com"

        # RSS 源列表（按可靠性排序）
        self.rss_sources = [
            {
                'url': 'https://www.coindesk.com/arc/outboundfeeds/rss/',
                'name': 'CoinDesk',
                'content_selectors': [
                    {'tag': 'div', 'attrs': {'class': re.compile(r'article-body|at-body|body-content')}},
                    {'tag': 'div', 'attrs': {'class': re.compile(r'content-body|entry-content')}},
                    {'tag': 'article'},
                ],
            },
            {
                'url': 'https://cointelegraph.com/rss',
                'name': 'CoinTelegraph',
                'content_selectors': [
                    {'tag': 'div', 'attrs': {'class': re.compile(r'post-content|article-content')}},
                    {'tag': 'article'},
                ],
            },
            {
                'url': 'https://cryptoslate.com/feed/',
                'name': 'CryptoSlate',
                'content_selectors': [
                    {'tag': 'div', 'attrs': {'class': re.compile(r'entry-content|post-content')}},
                    {'tag': 'article'},
                ],
            },
            {
                'url': 'https://cryptonews.com/feed/',
                'name': 'CryptoNews',
                'content_selectors': [
                    {'tag': 'div', 'attrs': {'class': re.compile(r'article-single__content|entry-content')}},
                    {'tag': 'article'},
                ],
            },
            {
                'url': 'https://bitcoinmagazine.com/feed',
                'name': 'BitcoinMagazine',
                'content_selectors': [
                    {'tag': 'div', 'attrs': {'class': re.compile(r'article-body|entry-content|m-detail--body')}},
                    {'tag': 'article'},
                ],
            },
        ]

        # 已见文章指纹集合（用于去重）
        self._seen_fingerprints = set()

    # --- 去重 ---

    def _fingerprint(self, title):
        """用标题生成指纹用于去重"""
        clean = re.sub(r'[^a-zA-Z0-9]', '', title.lower())
        return hashlib.md5(clean.encode()).hexdigest()

    def _is_duplicate(self, title):
        fp = self._fingerprint(title)
        if fp in self._seen_fingerprints:
            return True
        self._seen_fingerprints.add(fp)
        return False

    # --- RSS 获取 ---

    def _parse_published(self, entry):
        """从 RSS entry 中解析发布时间"""
        # 方法1: published_parsed
        if hasattr(entry, 'published_parsed') and entry.published_parsed:
            try:
                from time import mktime
                return datetime.fromtimestamp(mktime(entry.published_parsed), tz=timezone.utc)
            except Exception:
                pass
        # 方法2: published 字符串
        pub_str = entry.get('published') or entry.get('updated') or ''
        if pub_str:
            for fmt in ('%Y-%m-%dT%H:%M:%S%z', '%a, %d %b %Y %H:%M:%S %z',
                        '%a, %d %b %Y %H:%M:%S GMT', '%Y-%m-%dT%H:%M:%SZ'):
                try:
                    dt = datetime.strptime(pub_str.strip(), fmt)
                    if dt.tzinfo is None:
                        dt = dt.replace(tzinfo=timezone.utc)
                    return dt
                except ValueError:
                    continue
            # 尝试 ISO 格式
            try:
                clean = pub_str.strip()
                if clean.endswith('Z'):
                    clean = clean[:-1] + '+00:00'
                dt = datetime.fromisoformat(clean)
                if dt.tzinfo is None:
                    dt = dt.replace(tzinfo=timezone.utc)
                return dt
            except Exception:
                pass
        return None

    def get_articles_from_rss(self):
        """从所有 RSS 源获取文章列表，合并去重"""
        all_articles = []

        for source in self.rss_sources:
            rss_url = source['url']
            source_name = source['name']
            try:
                print(f"  正在获取 {source_name} RSS 源...")
                headers = NetworkHelper.get_headers(
                    accept='application/rss+xml,application/xml,text/xml,*/*'
                )
                # 使用较短超时：连接5秒，读取15秒
                resp = requests.get(rss_url, headers=headers, timeout=(5, 15))
                resp.raise_for_status()
                feed = feedparser.parse(resp.content)

                if not feed.entries:
                    print(f"    {source_name}: 无文章条目")
                    continue

                count = 0
                for entry in feed.entries[:20]:
                    title = entry.get('title', '').strip()
                    if not title:
                        continue
                    if self._is_duplicate(title):
                        continue

                    link = entry.get('link', '').strip()
                    # 提取摘要
                    summary = entry.get('summary', '') or entry.get('description', '')
                    if summary:
                        summary = re.sub(r'<[^>]+>', '', summary).strip()
                        if len(summary) > 500:
                            summary = summary[:500] + '...'

                    published = self._parse_published(entry)

                    all_articles.append({
                        'title': title,
                        'link': link,
                        'summary': summary,
                        'published': published,
                        'source': source_name,
                        'content': '',
                        'metrics': int(published.timestamp()) if published else 0,
                        '_content_selectors': source.get('content_selectors', []),
                    })
                    count += 1

                print(f"    {source_name}: 获取 {count} 篇文章")

            except KeyboardInterrupt:
                print(f"\n    用户中断，跳过剩余 RSS 源")
                break
            except Exception as e:
                print(f"    {source_name} 获取失败: {e}")
                continue

            time.sleep(1)

        print(f"\n  合计获取（去重后）: {len(all_articles)} 篇文章")
        return all_articles

    # --- 文章正文抓取 ---

    def get_article_content(self, url, content_selectors=None):
        """获取文章完整正文内容"""
        try:
            resp = NetworkHelper.get(url)
            soup = BeautifulSoup(resp.content, 'html.parser')

            # 提取发布时间
            published_dt = self._extract_published_from_page(soup)

            # 按站点特定选择器提取正文
            content = ''
            if content_selectors:
                for sel in content_selectors:
                    content_div = soup.find(sel['tag'], attrs=sel.get('attrs', {}))
                    if content_div:
                        content = self._extract_text(content_div)
                        if len(content) > 100:
                            break

            # 通用回退选择器
            if len(content) < 100:
                for selector in [
                    {'tag': 'div', 'attrs': {'class': re.compile(r'article.?body|post.?content|entry.?content|body.?content')}},
                    {'tag': 'article'},
                    {'tag': 'div', 'attrs': {'class': re.compile(r'content')}},
                ]:
                    el = soup.find(selector['tag'], attrs=selector.get('attrs', {}))
                    if el:
                        text = self._extract_text(el)
                        if len(text) > len(content):
                            content = text
                        if len(content) > 200:
                            break

            return content.strip(), published_dt

        except Exception as e:
            print(f"    获取文章内容失败 {url}: {e}")
            return "", None

    def _extract_text(self, element):
        """从 HTML 元素中提取干净的文本"""
        for tag in element.find_all(['script', 'style', 'nav', 'aside', 'footer',
                                      'iframe', 'form', 'button', 'svg']):
            tag.decompose()

        paragraphs = element.find_all('p')
        if paragraphs:
            texts = []
            for p in paragraphs:
                text = p.get_text(strip=True)
                if text and len(text) > 20:
                    texts.append(text)
            return '\n\n'.join(texts)
        else:
            return element.get_text(separator='\n', strip=True)

    def _extract_published_from_page(self, soup):
        """从网页 meta / time 标签提取发布时间"""
        try:
            for prop in ('article:published_time', 'datePublished', 'pubdate', 'date'):
                meta = soup.find('meta', property=prop) or soup.find('meta', attrs={'name': prop})
                if meta and meta.get('content'):
                    return self._parse_iso_datetime(meta['content'])

            for script in soup.find_all('script', type='application/ld+json'):
                try:
                    data = json.loads(script.string)
                    if isinstance(data, list):
                        data = data[0]
                    if isinstance(data, dict):
                        pub = data.get('datePublished')
                        if pub:
                            return self._parse_iso_datetime(pub)
                except Exception:
                    pass

            time_tag = soup.find('time', attrs={'datetime': True})
            if time_tag:
                return self._parse_iso_datetime(time_tag['datetime'])

        except Exception:
            pass
        return None

    @staticmethod
    def _parse_iso_datetime(s):
        """解析 ISO 格式时间字符串"""
        try:
            s = s.strip()
            if s.endswith('Z'):
                s = s[:-1] + '+00:00'
            dt = datetime.fromisoformat(s)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            return dt
        except Exception:
            return None


class Translator:
    """使用 Google Translate 非官方 API 进行翻译"""

    API_URL = "https://translate.googleapis.com/translate_a/single"

    def translate_to_chinese(self, text):
        """将英文翻译成中文，自动分块处理长文本"""
        if not text or not text.strip():
            return ""

        chunks = self._split_text(text, TRANSLATE_CHUNK)
        translated_parts = []

        for i, chunk in enumerate(chunks):
            translated = self._translate_chunk(chunk)
            translated_parts.append(translated)
            if i < len(chunks) - 1:
                time.sleep(1)

        return ''.join(translated_parts)

    def _translate_chunk(self, text):
        """翻译单个文本块"""
        try:
            params = {
                'client': 'gtx',
                'sl': 'en',
                'tl': 'zh-CN',
                'dt': 't',
                'q': text,
            }
            resp = NetworkHelper.get(
                self.API_URL,
                params=params,
                timeout=15,
                retries=2,
            )
            if resp.status_code == 200:
                result = resp.json()
                if result and result[0]:
                    return ''.join(item[0] for item in result[0] if item and item[0])
            return text
        except Exception as e:
            print(f"  翻译出错: {e}")
            return text

    @staticmethod
    def _split_text(text, max_len):
        """将文本按段落分块，每块不超过 max_len 字符"""
        if len(text) <= max_len:
            return [text]

        paragraphs = text.split('\n\n')
        chunks = []
        current = ''

        for para in paragraphs:
            if len(current) + len(para) + 2 > max_len:
                if current:
                    chunks.append(current)
                if len(para) > max_len:
                    sentences = re.split(r'(?<=[.!?])\s+', para)
                    sub = ''
                    for sent in sentences:
                        if len(sub) + len(sent) + 1 > max_len:
                            if sub:
                                chunks.append(sub)
                            sub = sent
                        else:
                            sub = sub + ' ' + sent if sub else sent
                    current = sub
                else:
                    current = para
            else:
                current = current + '\n\n' + para if current else para

        if current:
            chunks.append(current)
        return chunks


class WordDocumentCreator:
    """生成 Word 文档"""

    def __init__(self):
        self.doc = Document()
        self.translator = Translator()
        self._setup_styles()

    def _setup_styles(self):
        """设置文档默认样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.size = Pt(11)
        font.name = 'Calibri'
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15
        # 设置中文默认字体
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 设置文档属性
        core = self.doc.core_properties
        core.author = 'CryptoNews Collector'
        core.language = 'zh-CN'

    def add_title(self, title):
        """添加文档标题"""
        heading = self.doc.add_heading(title, 0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_meta_info(self, source_text):
        """添加文档元信息"""
        self.doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        self.doc.add_paragraph(f'文章来源: {source_text}')
        self.doc.add_paragraph('=' * 50)
        self.doc.add_paragraph()

    def add_article(self, article, index, translate=True):
        """添加一篇文章到文档"""
        title = article.get('title', '未知标题')
        link = article.get('link', '')
        summary = article.get('summary', '')
        content = article.get('content', '')
        source = article.get('source', '未知来源')

        # -- 文章标题 --
        if translate:
            translated_title = self.translator.translate_to_chinese(title)
            heading = self.doc.add_heading(f'文章 {index}: {translated_title}', level=1)
        else:
            heading = self.doc.add_heading(f'Article {index}: {title}', level=1)

        if heading.runs:
            heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        # -- 原文标题 --
        p = self.doc.add_paragraph()
        run = p.add_run('原文标题: ' if translate else 'Original Title: ')
        run.bold = True
        p.add_run(title)

        # -- 来源 --
        p = self.doc.add_paragraph()
        run = p.add_run('来源: ' if translate else 'Source: ')
        run.bold = True
        p.add_run(source)

        # -- 链接 --
        p = self.doc.add_paragraph()
        run = p.add_run('原文链接: ' if translate else 'Link: ')
        run.bold = True
        if link:
            self._add_hyperlink(p, link, link)
        else:
            p.add_run('N/A')

        # -- 发布时间 --
        if article.get('published'):
            try:
                pub_dt = article['published']
                beijing_tz = timezone(timedelta(hours=8))
                if pub_dt.tzinfo:
                    pub_dt = pub_dt.astimezone(beijing_tz)
                pub_text = pub_dt.strftime('%Y-%m-%d %H:%M:%S')
                p = self.doc.add_paragraph()
                run = p.add_run('发布时间: ' if translate else 'Published: ')
                run.bold = True
                p.add_run(f'{pub_text} (北京时间)')
            except Exception:
                pass

        # -- 摘要 --
        if summary:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run('【摘要 / Summary】' if translate else '【Summary】')
            run.bold = True
            # 摘要用略小字号区分
            summary_p = self.doc.add_paragraph(summary)
            for run in summary_p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(80, 80, 80)

            if translate:
                translated_summary = self.translator.translate_to_chinese(summary)
                p = self.doc.add_paragraph()
                p.add_run('【中文摘要】').bold = True
                self._add_chinese_paragraph(translated_summary)
                time.sleep(0.5)

        # -- 正文 --
        if content and len(content) > 80:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            p.add_run('【正文 / Full Text】' if translate else '【Full Text】').bold = True
            display_content = content if len(content) <= 3000 else content[:3000] + '\n\n[... 正文过长已截断 ...]'
            # 分段输出英文正文，避免单段过长导致格式失调
            for raw_para in display_content.split('\n\n'):
                raw_para = raw_para.strip()
                if raw_para:
                    self.doc.add_paragraph(raw_para)

            if translate:
                self.doc.add_paragraph()
                p = self.doc.add_paragraph()
                p.add_run('【中文翻译正文】').bold = True
                paragraphs = content.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para and len(para) > 10:
                        translated_para = self.translator.translate_to_chinese(para)
                        self._add_chinese_paragraph(translated_para)
                        time.sleep(0.8)
        elif content:
            # 正文太短，当作补充摘要显示
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            p.add_run('【补充内容】' if translate else '【Additional Content】').bold = True
            self.doc.add_paragraph(content)

        # -- 分隔线 --
        self.doc.add_paragraph()
        self.doc.add_paragraph('=' * 50)
        self.doc.add_paragraph()

    def _add_chinese_paragraph(self, text):
        """添加中文段落（带楷体字体设置）"""
        if not text or not text.strip():
            return
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        # 设置东亚字体为楷体
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '楷体')
        rFonts.set(qn('w:ascii'), 'Calibri')
        rFonts.set(qn('w:hAnsi'), 'Calibri')

    def _add_hyperlink(self, paragraph, url, text):
        """在段落中添加可点击超链接"""
        try:
            part = paragraph.part
            r_id = part.relate_to(
                url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                is_external=True
            )
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            c = OxmlElement('w:color')
            c.set(qn('w:val'), '0066CC')
            rPr.append(c)

            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)

            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')  # 10pt
            rPr.append(sz)

            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
        except Exception:
            # 回退为普通蓝色文字
            link_run = paragraph.add_run(text)
            link_run.font.color.rgb = RGBColor(0, 102, 204)
            link_run.font.size = Pt(10)

    def save(self, filename):
        """保存文档"""
        self.doc.save(filename)
        print(f"  文档已保存: {filename}")


def select_top_articles(articles, top_n=TOP_N):
    """
    选取最终文章:
    1. 优先选取今日(UTC)发布的文章
    2. 按发布时间倒序排列(最新优先)
    3. 如果今日不足 top_n 篇，用最近的文章补齐
    """
    now_utc = datetime.now(timezone.utc)
    today_utc = now_utc.date()

    todays = []
    others = []
    for a in articles:
        pub = a.get('published')
        if pub:
            try:
                pub_date = pub.date() if hasattr(pub, 'date') else None
                if pub_date == today_utc:
                    todays.append(a)
                    continue
            except Exception:
                pass
        others.append(a)

    todays.sort(key=lambda x: x.get('metrics', 0), reverse=True)
    others.sort(key=lambda x: x.get('metrics', 0), reverse=True)

    selected = todays[:top_n]
    if len(selected) < top_n:
        selected += others[: top_n - len(selected)]

    return selected


def _migrate_old_docx(base_dir, history_dir):
    """将根目录中残留的 .docx 文件迁移到 history/ 子目录"""
    moved = 0
    for fname in os.listdir(base_dir):
        if fname.lower().endswith('.docx'):
            src = os.path.join(base_dir, fname)
            dst = os.path.join(history_dir, fname)
            if os.path.isfile(src):
                try:
                    shutil.move(src, dst)
                    moved += 1
                except Exception as e:
                    print(f"  迁移文件失败 {fname}: {e}")
    if moved:
        print(f"  已将 {moved} 个历史文档迁移到 history/ 目录")


def main():
    print("=" * 60)
    print("  加密货币新闻文章收集器")
    print("  " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    print("=" * 60)

    # -- 解析命令行参数 --
    translate = True
    generate_english = False

    if len(sys.argv) > 1:
        arg = sys.argv[1].lower()
        if arg in ('--english', '-e'):
            translate = False
            generate_english = True
            print("模式: 仅生成英文版")
        elif arg in ('--both', '-b'):
            translate = True
            generate_english = True
            print("模式: 同时生成英文版 + 中文翻译版")
        elif arg in ('--help', '-h'):
            print("\n用法: python coindesk_collector.py [选项]")
            print("  (无参数)          仅生成中文翻译版")
            print("  --english, -e     仅生成英文版（快速，不翻译）")
            print("  --both, -b        同时生成英文版和中文翻译版")
            print("  --help, -h        显示此帮助信息")
            return
        else:
            print(f"未知参数: {sys.argv[1]}，使用 --help 查看帮助")
            return
    else:
        print("模式: 生成中文翻译版（默认）")

    # -- 初始化收集器 --
    collector = CoinDeskCollector()

    # -- 获取文章列表 --
    print("\n正在从 RSS 源获取文章...")
    articles = collector.get_articles_from_rss()

    if not articles:
        print("\n未能从任何来源获取到文章，请检查网络连接后重试。")
        return

    # -- 候选文章 --
    articles.sort(key=lambda x: x.get('metrics', 0), reverse=True)
    candidates = articles[:CANDIDATE_N]

    print(f"\n正在获取 {len(candidates)} 篇候选文章的完整内容...")
    for i, article in enumerate(candidates, 1):
        print(f"  [{i}/{len(candidates)}] {article['title'][:60]}...")
        try:
            if not article.get('content'):
                selectors = article.pop('_content_selectors', [])
                content, page_published = collector.get_article_content(
                    article['link'], content_selectors=selectors
                )
                article['content'] = content
                if page_published and not article.get('published'):
                    article['published'] = page_published
                if content:
                    print(f"    获取到 {len(content)} 字符正文")
                else:
                    print(f"    未能获取正文（将仅使用摘要）")
        except KeyboardInterrupt:
            print(f"\n  用户中断，使用已获取的 {i-1} 篇文章内容继续...")
            break
        except Exception as e:
            print(f"    获取失败: {e}")
        time.sleep(REQUEST_DELAY)

    # -- 选择最终文章 --
    top_articles = select_top_articles(candidates, TOP_N)

    if not top_articles:
        print("\n筛选后无可用文章。")
        return

    print(f"\n最终选择 {len(top_articles)} 篇文章:")
    for i, a in enumerate(top_articles, 1):
        pub_str = ''
        if a.get('published'):
            try:
                pub_str = a['published'].strftime(' (%Y-%m-%d %H:%M)')
            except Exception:
                pass
        print(f"  {i}. [{a.get('source', '?')}] {a['title'][:60]}{pub_str}")

    # -- 收集所有来源名 --
    sources = sorted(set(a.get('source', '') for a in top_articles if a.get('source')))
    source_text = ' & '.join(sources) if sources else 'Crypto News'

    today = datetime.now().strftime('%Y-%m-%d')
    base_dir = os.path.dirname(os.path.abspath(__file__))
    history_dir = os.path.join(base_dir, 'history')
    os.makedirs(history_dir, exist_ok=True)

    # -- 自动迁移根目录中的旧 docx 文件到 history/ --
    _migrate_old_docx(base_dir, history_dir)

    # -- 生成英文版 --
    if generate_english or not translate:
        print("\n正在生成英文版文档...")
        doc_en = WordDocumentCreator()
        doc_en.add_title(f'Crypto News - {today}')
        doc_en.add_meta_info(source_text)

        for i, article in enumerate(top_articles, 1):
            doc_en.add_article(article, i, translate=False)

        filename_en = os.path.join(history_dir, f'CryptoNews_English_{today}.docx')
        doc_en.save(filename_en)

    # -- 生成中文翻译版 --
    if translate:
        print("\n正在翻译并生成中文版文档（可能需要几分钟）...")
        doc_cn = WordDocumentCreator()
        doc_cn.add_title(f'加密货币热门文章 - {today}')
        doc_cn.add_meta_info(source_text)

        for i, article in enumerate(top_articles, 1):
            print(f"  翻译第 {i}/{len(top_articles)} 篇: {article['title'][:50]}...")
            doc_cn.add_article(article, i, translate=True)

        filename_cn = os.path.join(history_dir, f'CryptoNews_中文_{today}.docx')
        doc_cn.save(filename_cn)

    print("\n" + "=" * 60)
    print("  处理完成！")
    print("=" * 60)


if __name__ == "__main__":
    main()
